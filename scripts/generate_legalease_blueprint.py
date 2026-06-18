#!/usr/bin/env python3
"""
Generate LegalEase full system blueprint as Word (.docx) and PDF.
Run: py scripts/generate_legalease_blueprint.py
Output: docs/exports/LEGALEASE_BLUEPRINT.docx + LEGALEASE_BLUEPRINT.pdf
"""
from __future__ import annotations

import textwrap
from datetime import datetime, timezone
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "exports"
OUT_DOCX = OUT_DIR / "LEGALEASE_BLUEPRINT.docx"
OUT_PDF = OUT_DIR / "LEGALEASE_BLUEPRINT.pdf"

VERSION = "3.0"
GENERATED = datetime.now(timezone.utc).strftime("%Y-%m-%d %H:%M UTC")


# ---------------------------------------------------------------------------
# Blueprint content — sections with optional diagram (ASCII)
# ---------------------------------------------------------------------------

SECTIONS: list[dict] = [
    {
        "title": "1. Executive Summary",
        "body": """
LegalEase is an Indian legal AI assistant for law firms and solo practitioners.
It combines local document RAG (your PDFs), live web legal search (Gemini),
deep hybrid research reports, and a continuous self-improvement loop.

Core design rule: Knowledge Base answers come ONLY from uploaded documents +
local Ollama/LM Studio. Gemini never writes KB answers.

Stack: Next.js frontend | FastAPI backend | FAISS vector store | Ollama/LM Studio
local LLM | Gemini API (Open Law + Hybrid fusion + Settings coach only).

Deployment model: Single-lawyer self-hosted (production-ready). Multi-tenant cloud
requires additional per-user model routing (not required for solo use).
""",
    },
    {
        "title": "2. Master System Architecture",
        "diagram": """
+------------------------------------------------------------------+
|                        LEGALEASE v3.0                             |
+------------------------------------------------------------------+
|  FRONTEND (Next.js)          |  BACKEND (FastAPI)                 |
|  / chat, documents, matters  |  /api/v1/* REST + SSE stream       |
|  settings, analytics, CRM    |  chat_service, mode_router         |
+--------------+---------------+---------------+--------------------+
               |                               |
               v                               v
+--------------+--------+            +---------+---------+
|  SQLite DB   | FAISS  |            |  Ollama / LM Studio|
|  users, CRM, | per    |            |  KB synthesis ONLY |
|  billing,    | user / |            +---------+---------+
|  learning    | matter |                      |
+--------------+--------+                      |
               |                               |
               v                               v
+--------------+--------+            +---------+---------+
|  Data/       | OCR    |            |  Gemini API        |
|  uploads,    | EasyOCR|            |  Open Law web only |
|  exports     |        |            |  Hybrid fusion     |
+--------------+--------+            |  Settings coach    |
                                     +--------------------+
""",
        "body": """
Data flow: User -> Next.js -> FastAPI -> Mode Router -> (KB | Open Law | Hybrid)
-> Local LLM and/or Gemini -> Response -> Learning hooks -> Background automation.

Key entry files:
  backend/app/main.py          API startup, middleware, schedulers
  backend/app/services/chat_service.py   Chat orchestration
  backend/app/services/mode_router.py    Mode selection
  rag.py / kb_pipeline.py      RAG retrieval pipeline
  llms.py                      Ollama / LM Studio client
  web/app/(app)/page.tsx       Main chat UI
""",
    },
    {
        "title": "3. Three Chat Modes — Routing Diagram",
        "diagram": """
                    [ User sends message ]
                              |
                              v
                    [ Parse legal intent ]
                              |
                              v
                    [ Merge follow-up context ]
                              |
                              v
                       +------+------+
                       | Mode Router |
                       +------+------+
            +-------------+-------------+-------------+
            |             |             |             |
            v             v             v             v
    [Knowledge Base] [Open Law]   [Hybrid/Juris]  [Overrides]
            |             |             |        empty KB->web
            v             v             v        case in KB->hybrid
    FAISS + Ollama   Gemini +      KB + Web +
    local LLM only   Google Search  Gemini fusion
            |             |             |
            +-------------+-------------+
                              |
                              v
                    [ Response to user ]
                              |
                              v
                    [ Feedback + learning hooks ]
""",
        "body": """
| Mode            | UI Name        | Engine              | Best for                    |
|-----------------|----------------|---------------------|-----------------------------|
| knowledge_base  | Knowledge Base | FAISS + Ollama      | Your uploaded documents     |
| web_search      | Open Law       | Gemini + web search | Live statutes, cases, news  |
| deep_case       | Hybrid (Pro+)  | KB + Gemini fusion  | Full research reports       |

Smart overrides (mode_router.py):
  - Empty KB + case question -> may suggest Open Law
  - KB mode + case explanation -> may upgrade to Hybrid (Pro membership)
""",
    },
    {
        "title": "4. Knowledge Base Engine — Feature Diagram",
        "diagram": """
  [Your Question] -----> [Query Understanding / Intent]
              |                    |
              |                    v
              |           [Memory Injection: persona, facts]
              |                    |
              v                    v
  [Matter Scope?] -----> [Query Expansion]
              |                    |
              +--------> [FAISS Dense Search]
              |                    |
              +--------> [Keyword / Sparse Search]
                                   |
                                   v
                          [Cross-Encoder Rerank]
                                   |
                                   v
                          [Confidence Gate]
                          /              \\
                    LOW score          HIGH score
                        |                  |
                        v                  v
                   [NOT_FOUND]    [Ollama Synthesis]
                        |                  |
                        +--------+---------+
                                 |
                                 v
                          [Answer + Citations]
                                 |
                                 v
                          [Learning hook log]
""",
        "body": """
Steps:
  1. Query understanding — section lookup, comparison, list, case intent
  2. Memory injection — persona, facts, coach style (never legal substance from Gemini)
  3. Hybrid retrieval — dense FAISS + sparse keyword + optional cross-encoder rerank
  4. Confidence gate — low scores return NOT_FOUND instead of hallucination
  5. Synthesis — local LLM writes answer from chunks only (GEMINI_KB_SYNTHESIS=0)
  6. Per-matter scoping — matter_id limits search to that case file's documents

Indexing: Upload PDF/image -> OCR if scanned -> chunk -> embed -> FAISS index
Path: faiss_indexes/user_{id}/matter_{matter_id}/

Key files: kb_pipeline.py, rag.py, answer_orchestrator.py, llms.py, matter_index.py
""",
    },
    {
        "title": "5. Open Law (Web Intelligence) — Feature Diagram",
        "diagram": """
  [User: Open Law question]
            |
            v
  [Filter KB/Hybrid history]  <-- prevents document leak
            |
            v
  [Classify depth: quick | standard | detailed | comparison]
            |
            v
  [Build self-contained query]
            |
            v
  [Gemini + Google Search Grounding]
            |
            v
  [KB-leak detection] ---- retry without history if doc dump detected
            |
            v
  [Format: bullets | table | short fact]
            |
            v
  [Response + source URLs + trust badges]
""",
        "body": """
Adaptive response sizing (web_intelligence.py):
  - Quick fact ("Who is CJI?") -> 2-4 sentences
  - Standard question -> concise bullets (~260 words)
  - "Explain in detail" -> structured sections (~480 words)
  - Comparison ("IPC 300 vs 307") -> markdown comparison table

Safety: No KB history sent to Gemini. No "Gemini" branding in UI.
Gemini daily quota enforced by plan (Free 15 / Pro 200 / Legal Pro 1000).
""",
    },
    {
        "title": "6. Hybrid / Jurisprudence Engine — Feature Diagram",
        "diagram": """
  [Hybrid question]
        |
        v
  [Parallel fetch]
   /            \\
  v              v
[KB RAG         [Open Law
 14 chunks]      web search]
  |              |
  +------+-------+
         |
         v
  [Gemini Jurisprudence Fusion]
         |
         v
  +------+------+------+------+------+------+------+
  | Exec | Doc  | Pub  | Stat | Case | Recon| Rec  |
  | Summ | Intel| Intel| Frame| Law  | KB/Web|    |
  +------+------+------+------+------+------+------+
         |
         v
  [Export: DOCX / PDF / Markdown (client-safe option)]
""",
        "body": """
Conflict rule: Uploaded documents win when KB and web disagree on sections/statutes.
Requires Pro or Legal Pro membership (Free users downgraded to KB).
Includes contradiction checker and citation verifier on output.
""",
    },
    {
        "title": "7. Gemini vs Ollama — Separation of Duties",
        "diagram": """
  +---------------------------+       +---------------------------+
  |     GEMINI (cloud)        |       |   OLLAMA / LM STUDIO      |
  +---------------------------+       +---------------------------+
  | Open Law web answers      |       | KB answer synthesis       |
  | Hybrid fusion reports     |  X--> | Reads YOUR documents only |
  | Settings coach (meta)     |       | Never uses Gemini for KB  |
  +---------------------------+       +---------------------------+
         |                                       ^
         | style/format ONLY                     |
         +---------------------------------------+
              (coach never injects legal substance)

  Env locks: GEMINI_KB_SYNTHESIS=0  |  GEMINI_OLLAMA_TUNING=1
""",
        "body": """
| Task                         | Engine        |
|------------------------------|---------------|
| Answer from uploaded PDFs    | Ollama only   |
| Live web legal search        | Gemini only   |
| Hybrid research report       | Both          |
| Tune Ollama style/retrieval  | Gemini coach  |
| Answer legal Q in coach      | FORBIDDEN     |
""",
    },
    {
        "title": "8. AI Memory & Persona",
        "diagram": """
  [Settings: AI Memory]
        |
        +-- Persona (warm / professional / concise / detailed)
        +-- Practice area (e.g. Criminal litigation)
        +-- Facts (key-value, manual + auto-suggested)
        +-- Communication notes
        +-- Coach memories (style only)
        +-- Thread summaries
        |
        v
  [Injected every chat turn]
        |
        v
  [KB + Open Law + Hybrid prompts]
""",
        "body": """
Auto-facts appear in amber badges — user can edit/delete wrong ones.
Past chat RAG: semantic search over prior threads ("what did we conclude about...").
API: /api/v1/memory — profile, facts CRUD, reindex chats.
Key file: backend/app/core/user_memory.py
""",
    },
    {
        "title": "9. Learning & Feedback Loop",
        "diagram": """
  [Every chat turn logged: mode, query, preview, chunks, found/not found]
        |
        v
  [User feedback: thumbs up / thumbs down + comment]
        |
        +---> Adaptive learning (chunk boosts, query expansions)
        +---> Neural training pairs (thumbs-up Q->passage)
        +---> Answer memory (instant replay on similar Q)
        +---> Gemini coach (on thumbs-down with comment)
        |
        v
  [Better retrieval + style on next similar question]
""",
        "body": """
Thumbs up: boosts chunks, adds neural pair, answer memory, triggers automation pipeline.
Thumbs down + comment: records complaint, triggers coach analysis if enabled.
Mode stats track hit rate per KB / Open Law / Hybrid.
Key files: adaptive_learning.py, learning_engine.py
""",
    },
    {
        "title": "10. Gemini Ollama Coach (Settings Only)",
        "diagram": """
  TRIGGERS (Settings only — NEVER during normal chat):
    [Run tuning cycle] [Save directives] [Thumbs-down comment]
    [Daily auto-scheduler if new feedback]
        |
        v
  [Analyze feedback patterns + user instructions]
        |
        v
  +-----+-----+-----+-----+
  |Persona|Style|Search|Pairs|
  |suggest|facts|phrase|from|
  |       |     |improve|YOUR|
  |       |     |       |thumbs|
  +-----+-----+-----+-----+
        |
        v
  ANTI-BIAS GUARDS block legal substance, banned phrases, Gemini Q->A pairs
""",
        "body": """
Coach outputs ALLOWED: persona suggestion, style facts, search phrasing, verified training pairs.
Coach outputs FORBIDDEN: legal holdings, IPC/BNS text, Gemini-generated training pairs.
Schedule defaults: 1 day interval, 1 min new feedback, check every 30 min.
Key files: gemini_ollama_coach.py, coach_scheduler.py
""",
    },
    {
        "title": "11. Auto-Improvement Pipeline (Fully Automated)",
        "diagram": """
  [Thumbs up / down / helpful / copy]
        |
        v
  [schedule_improvement_pipeline] (background thread, deduped per user)
        |
        v
  +-----+-----+-----+-----+-----+
  |Neural|Re-  |Export|Ollama|Switch|
  |train |index|Model-|create|to    |
  |embed |KB   |file  |legale|tuned |
  |      |     |      |ase   |model |
  +-----+-----+-----+-----+-----+
        |
        v
  [At 20+ thumbs-up: auto ollama create legalease-tuned -f Modelfile]
        |
        v
  [Settings panel shows progress: X/20 thumbs-up]
""",
        "body": """
Env vars: IMPROVEMENT_AUTO=1, OLLAMA_AUTO_REINDEX=1, OLLAMA_AUTO_CREATE=1,
OLLAMA_AUTO_USE_TUNED=1, OLLAMA_AUTO_EXPORT_MIN_THUMBS=20

Export path: Data/ollama_exports/{user_id}/{timestamp}/
Files: Modelfile, training.jsonl, README.txt
Log: Data/ollama_exports/automation_log.jsonl
Key file: backend/app/core/improvement_automation.py
""",
    },
    {
        "title": "12. Neural Training & Modelfile Export",
        "diagram": """
  LEVEL 1 — Embedding fine-tune:
    [Thumbs-up pairs] -> [SentenceTransformer train] -> [Re-index FAISS]

  LEVEL 2 — Modelfile export:
    [20+ thumbs-up] -> [Export Modelfile + JSONL] -> [ollama create legalease-tuned]

  Both compound over time for better retrieval AND better answer style.
""",
        "body": """
Neural: NEURAL_FINETUNE_AUTO=1, min 4 pairs, trains MiniLM variant on Q->passage pairs.
Modelfile: base model + system prompt + few-shot examples from verified feedback.
After embedding train: MUST re-index documents for new embeddings to take effect.
""",
    },
    {
        "title": "13. Documents & Knowledge Base Health",
        "diagram": """
  [/documents page]
        |
        +-- Drag-drop upload (PDF, images)
        +-- Assign to matter
        +-- OCR for scans (OCR_ENABLED=1)
        +-- Entity extraction + timeline
        +-- Delete / reindex
        |
        v
  [KbHealthPanel: embeddings online, FAISS count, stale warnings]
        |
        v
  [Auto-reindex if index empty but docs exist]
""",
        "body": """
API: /api/v1/documents — upload, list, index, kb/reindex-auto, entities, timeline.
Per-matter upload: ?matter_id= scopes index to that case file.
""",
    },
    {
        "title": "14. Matters & Case Scoping",
        "diagram": """
  [Create Matter / Case File]
        |
        +-- Link documents
        +-- Notes
        +-- Client portal token link
        +-- E-sign mock flow
        |
        v
  [Chat: select Case File before KB/Hybrid question]
        |
        v
  [FAISS searches ONLY that matter's index]
""",
        "body": """
Index path: faiss_indexes/user_{id}/matter_{matter_id}/
Unlinked docs: faiss_indexes/user_{id}/_unlinked/
Matter autopilot: entity extraction + suggested research queries.
API: /api/v1/matters
""",
    },
    {
        "title": "15. Practice SaaS Modules",
        "diagram": """
  +----------+ +----------+ +----------+ +----------+
  | Dashboard| | Billing  | | Intake   | |Discovery |
  | /dashboard| |/billing | | /intake  | |/discovery|
  +----------+ +----------+ +----------+ +----------+
  +----------+ +----------+ +----------+ +----------+
  | Premium  | | Tools    | | Drafting | | Portal   |
  | /premium | | /tools   | |/drafting | |/portal/* |
  +----------+ +----------+ +----------+ +----------+
        |
        v
  [All user-scoped via JWT user_id in SQLite]
""",
        "body": """
| Module    | Route       | Purpose                          |
|-----------|-------------|----------------------------------|
| Dashboard | /dashboard  | Practice overview, stats           |
| Billing   | /billing    | Time entries, GST invoices       |
| Intake    | /intake     | CRM lead capture                 |
| Discovery | /discovery  | E-discovery triage + job queue   |
| Premium   | /premium    | Witness sim, BNS audit, deal rooms|
| Tools     | /tools      | IPC lookup, fees, contract tools |
| Drafting  | /drafting   | Template drafting studio         |
| Portal    | /portal/*   | Client-facing matter view        |
| E-sign    | /esign/mock | Mock signing flow                |
""",
    },
    {
        "title": "16. Premium Tools Detail",
        "diagram": """
  [Premium /premium]
     |
     +-- Witness / Mock Trial Simulator (disposition: evasive, hostile...)
     +-- Precedent Tree (citation graph + judge analytics)
     +-- BNS Auditor (IPC->BNS risk scan on text/docs)
     +-- Deal Rooms (multi-doc diligence + anomaly detection)
     +-- PII Redaction (detect, whitelist, redact by entity type)
     +-- Redline Engine (clause revision with diff HTML)
""",
        "body": """
API: /api/v1/premium — witness sessions, BNS audit, deal rooms, PII, redline.
Separate premium learning store for tool feedback.
""",
    },
    {
        "title": "17. Billing, CRM & Trust",
        "diagram": """
  BILLING (/billing):
    [Time entry] -> [AI narrative polish] -> [Invoice + GST] -> [Summary]

  CRM (/intake):
    [Lead] -> [AI intent classify] -> [Pipeline stages] -> [Follow-up email draft]
    [Public intake form] -> unauthenticated lead capture

  TRUST (/billing ledger):
    [Matter trust balance] <-> [Transaction ledger]
""",
        "body": """
Billing API: /api/v1/billing — time entries, invoices, narrative polish, lexicon.
CRM API: /api/v1/crm — leads, classification, corrections, follow-ups.
Trust API: /api/v1/trust — matter trust balances and transactions.
Subscription billing: mock upgrade (Stripe not wired); practice billing is fully implemented.
""",
    },
    {
        "title": "18. E-Discovery & Research",
        "diagram": """
  [Upload batch to matter]
        |
        v
  [Background job queue (Redis/SQLite)]
        |
        v
  [Triage: relevance / privilege-style scoring]
        |
        v
  [Search within batch + tag/classify items]
        |
        v
  [Job status polling in /discovery UI]
""",
        "body": """
API: /api/v1/ediscovery — batches, jobs, triage, search, review.
Worker: scripts/ediscovery_worker.py (optional background process).
Research log: /api/v1/research — query expansion, history, feedback.
""",
    },
    {
        "title": "19. Settings, Engine Status & Analytics",
        "diagram": """
  SETTINGS (/settings):
    AI Memory | Coach | Auto-improvement | Profile | LLM test | OCR status

  ENGINE STATUS BAR (chat top):
    [KB chips] [Web Intel] [LLM online] [Memory active] [Gemini quota]
    Polls every ~45 seconds

  ANALYTICS (/analytics):
    Mode stats | Learned queries | Deal rooms | Witness sessions | JSONL export
""",
        "body": """
Settings sections: persona, facts, neural train, Ollama coach, auto-improvement pipeline,
daily auto-coaching toggle, Modelfile export, LLM backend status.
Analytics API: /api/v1/learning/analytics/full
""",
    },
    {
        "title": "20. Auth, Membership & Rate Limits",
        "diagram": """
  [Login / Register] -> [JWT Bearer token]
        |
        v
  [All /api/v1/* routes require auth]
        |
        +-- Free: 2 doc upload cap, Hybrid blocked, Gemini 15/day
        +-- Pro: Hybrid allowed, Gemini 200/day, unlimited docs
        +-- Legal Pro: Gemini 1000/day
        |
        v
  [Rate limit: 120 req/min general, 40 req/min on /chat]
""",
        "body": """
Auth: backend/app/core/auth.py, legacy_saas/legalease_auth.py
Rate limits: backend/app/middleware/rate_limit.py (in-memory, per IP/token prefix)
Membership stored in JWT + users table.
""",
    },
    {
        "title": "21. Frontend Pages Map",
        "body": """
/login          Authentication
/               Main AI chat (KB / Open Law / Hybrid)
/documents      Upload & index knowledge base
/matters        Case files
/settings       Memory, coach, LLM config, auto-improvement
/analytics      Learning & usage stats
/billing        Time & billing + trust ledger
/premium        Advanced tools
/tools          Legal calculators & lookups
/discovery      E-discovery
/intake         CRM
/drafting       Document drafting studio
/dashboard      Practice dashboard
/portal/[token] Client portal (read-only matter view)
/esign/mock/*   Mock e-signature flow
""",
    },
    {
        "title": "22. API Reference (All Groups)",
        "body": """
Base URL: /api/v1

| Prefix              | Purpose                                    |
|---------------------|--------------------------------------------|
| /chat               | Send message, stream SSE, export report    |
| /learning           | Feedback, stats, neural, coach, automation |
| /memory             | Profile, facts, reindex chats              |
| /documents          | Upload, list, KB health, reindex           |
| /engines            | Status bar, watchlist, autopilot             |
| /matters            | Case CRUD, notes, docs                       |
| /sessions           | Thread history, attachments                  |
| /premium            | Witness, BNS audit, deal rooms, PII          |
| /billing            | Time entries, invoices, narratives           |
| /crm                | Leads, intake classification                 |
| /trust              | Trust ledger                                 |
| /ediscovery         | Review batches, jobs                         |
| /research           | Research log                                 |
| /templates, /clauses| Drafting templates and clause library        |
| /portal             | Client portal                                |
| /esign              | E-signature requests                         |
| /practice           | Dashboard overview, public intake            |
| /health             | Live probe (no ML import)                    |
""",
    },
    {
        "title": "23. Environment Variables (Essential)",
        "body": """
# Local KB brain
LLM_BACKEND=ollama
OLLAMA_MODEL=gemma3:4b
OLLAMA_BASE_URL=http://127.0.0.1:11434

# Web intel + coach
GEMINI_API_KEY=your_key
GEMINI_OLLAMA_TUNING=1
GEMINI_KB_SYNTHESIS=0          # MUST stay 0

# Learning & automation
LEARNING_ENGINE_ENABLED=1
NEURAL_FINETUNE_ENABLED=1
NEURAL_FINETUNE_AUTO=1
IMPROVEMENT_AUTO=1
OLLAMA_AUTO_REINDEX=1
OLLAMA_AUTO_CREATE=1
OLLAMA_AUTO_USE_TUNED=1
OLLAMA_AUTO_EXPORT_MIN_THUMBS=20
OLLAMA_TUNED_MODEL_NAME=legalease-tuned

# Daily coach
COACH_AUTO_SCHEDULE=1
COACH_AUTO_INTERVAL_DAYS=1
COACH_AUTO_MIN_NEW_FEEDBACK=1
COACH_AUTO_CHECK_INTERVAL_SEC=1800

See .env.example for full RAG, OCR, rate limit, and Redis settings.
""",
    },
    {
        "title": "24. Continuous Improvement Loop (Master)",
        "diagram": """
  DAILY USE                    AUTOMATIC (no Gemini in chat)
  ---------                    -----------------------------
  Chat KB/Open Law      -->    Adaptive chunk boosts
  Thumbs up/down        -->    Answer memory replay
  Settings instructions -->    Query expansion learning
                               Neural embedding train
                               KB auto re-index

  SETTINGS COACH (Gemini meta only)
  ---------------------------------
  Analyze feedback --> Apply style --> Collect verified pairs
                    --> Export Modelfile --> ollama create
                    --> Better answers on next similar question

  Four compounding layers:
    1. Immediate  — answer memory + chunk boosts
    2. Short-term — query expansions + persona/style
    3. Medium-term — neural embedding fine-tune
    4. Long-term  — Modelfile -> custom Ollama model
""",
        "body": """
Gemini's role is coach only — never answers KB questions, cannot inject legal substance.
Automation runs in background threads; never blocks chat.
""",
    },
    {
        "title": "25. Project File Structure",
        "body": """
Legal_AI_Final 3/
  backend/app/           FastAPI application
    api/v1/endpoints/    REST route handlers (19 modules)
    core/                Business logic (50+ modules)
    services/            Chat, hybrid, mode routing
  web/                   Next.js frontend
    app/(app)/           Main pages (chat, documents, settings...)
    components/          React UI components
    lib/api.ts           API client
  rag.py                 Core RAG retrieval
  kb_pipeline.py         Multi-stage KB pipeline
  llms.py                Ollama / LM Studio client
  app.py                 Legacy Streamlit monolith
  legacy_saas/           Auth, legacy API routes
  Data/                  Uploads, FAISS indexes, exports
  docs/                  Documentation + this blueprint export
  scripts/               Workers, eval, blueprint generator
  tests/                 Pytest suites (phases 1-4, SaaS, automation)
  deploy/                Docker, nginx TLS config
""",
    },
    {
        "title": "26. Quick Start Checklist",
        "body": """
[ ] Ollama or LM Studio running locally
[ ] GEMINI_API_KEY in root .env
[ ] GEMINI_OLLAMA_TUNING=1 and GEMINI_KB_SYNTHESIS=0
[ ] LLM_BACKEND=ollama (for tuned model support)
[ ] ollama CLI on PATH (for auto-create)
[ ] Upload documents on /documents
[ ] Enable coach in Settings -> AI memory
[ ] Chat and use thumbs-up/down with comments
[ ] At 20+ thumbs-up: auto Modelfile export + ollama create legalease-tuned
[ ] Re-index after neural embedding training (automatic when IMPROVEMENT_AUTO=1)
[ ] Optional: run ediscovery worker for background discovery jobs
""",
    },
]


def _add_diagram_paragraph(doc, text: str, style_name: str = "No Spacing") -> None:
    from docx.shared import Pt
    from docx.enum.text import WD_LINE_SPACING

    p = doc.add_paragraph(style=style_name)
    run = p.add_run(text.strip())
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)


def build_docx() -> Path:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.85)
        section.right_margin = Inches(0.85)

    # Cover
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("LegalEase AI Platform")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run(f"Complete System Blueprint & Feature Diagrams\nVersion {VERSION}")
    r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x44, 0x55, 0x66)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta.add_run(f"Generated: {GENERATED}\nIndian Legal AI · Next.js + FastAPI + Ollama + Gemini")
    r3.font.size = Pt(11)

    doc.add_page_break()

    # TOC
    doc.add_heading("Table of Contents", level=1)
    for sec in SECTIONS:
        doc.add_paragraph(sec["title"], style="List Number")
    doc.add_page_break()

    # Sections
    for sec in SECTIONS:
        doc.add_heading(sec["title"], level=1)
        if sec.get("diagram"):
            doc.add_heading("Diagram", level=2)
            _add_diagram_paragraph(doc, sec["diagram"])
        if sec.get("body"):
            doc.add_heading("Details", level=2)
            for block in sec["body"].strip().split("\n\n"):
                block = block.strip()
                if not block:
                    continue
                if block.startswith("|") or block.startswith("["):
                    _add_diagram_paragraph(doc, block)
                else:
                    for line in block.split("\n"):
                        doc.add_paragraph(line.strip())
        doc.add_page_break()

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    return OUT_DOCX


def _ascii_safe(text: str) -> str:
    """Replace unicode chars that break Helvetica PDF encoding."""
    replacements = {
        "\u2014": "-",
        "\u2013": "-",
        "\u2192": "->",
        "\u2190": "<-",
        "\u2022": "*",
        "\u00b7": ".",
        "\u2194": "<->",
        "\u2713": "[x]",
        "\u2714": "[x]",
        "\u2717": "[ ]",
        "\u2718": "[ ]",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text.encode("ascii", errors="replace").decode("ascii")


def build_pdf() -> Path:
    from fpdf import FPDF

    class BlueprintPDF(FPDF):
        def header(self):
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(100, 100, 100)
            self.cell(0, 8, _ascii_safe(f"LegalEase Blueprint v{VERSION}"), align="L")
            self.ln(4)

        def footer(self):
            self.set_y(-15)
            self.set_font("Helvetica", "I", 8)
            self.set_text_color(120, 120, 120)
            self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

        def section_title(self, title: str):
            title = _ascii_safe(title)
            self.set_font("Helvetica", "B", 14)
            self.set_text_color(30, 58, 95)
            self.multi_cell(0, 8, title)
            self.ln(2)
            self.set_draw_color(30, 58, 95)
            self.line(10, self.get_y(), 200, self.get_y())
            self.ln(4)

        def sub_title(self, title: str):
            title = _ascii_safe(title)
            self.set_font("Helvetica", "B", 11)
            self.set_text_color(60, 60, 60)
            self.multi_cell(0, 6, title)
            self.ln(2)

        def body_text(self, text: str):
            text = _ascii_safe(text)
            self.set_font("Helvetica", "", 10)
            self.set_text_color(30, 30, 30)
            self.multi_cell(0, 5, text)
            self.ln(2)

        def diagram_text(self, text: str):
            text = _ascii_safe(text)
            self.set_font("Courier", "", 7)
            self.set_text_color(20, 20, 20)
            self.set_fill_color(245, 247, 250)
            x = self.get_x()
            y = self.get_y()
            lines = text.strip().split("\n")
            h = len(lines) * 3.5 + 4
            if y + h > 270:
                self.add_page()
                y = self.get_y()
            self.rect(10, y, 190, h, style="F")
            self.set_xy(12, y + 2)
            for line in lines:
                self.cell(0, 3.5, line[:110], ln=True)
            self.ln(4)

    pdf = BlueprintPDF()
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()

    # Cover
    pdf.ln(40)
    pdf.set_font("Helvetica", "B", 26)
    pdf.set_text_color(30, 58, 95)
    pdf.cell(0, 14, "LegalEase AI Platform", align="C", ln=True)
    pdf.set_font("Helvetica", "", 16)
    pdf.set_text_color(68, 85, 102)
    pdf.cell(0, 10, "Complete System Blueprint", align="C", ln=True)
    pdf.cell(0, 10, f"and Feature Diagrams  |  v{VERSION}", align="C", ln=True)
    pdf.ln(10)
    pdf.set_font("Helvetica", "I", 11)
    pdf.cell(0, 8, f"Generated: {GENERATED}", align="C", ln=True)
    pdf.cell(0, 8, "Indian Legal AI | Next.js + FastAPI + Ollama + Gemini", align="C", ln=True)

    # TOC
    pdf.add_page()
    pdf.section_title("Table of Contents")
    for i, sec in enumerate(SECTIONS, 1):
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 6, _ascii_safe(f"{i}. {sec['title']}"), ln=True)

    # Content
    for sec in SECTIONS:
        pdf.add_page()
        pdf.section_title(sec["title"])
        if sec.get("diagram"):
            pdf.sub_title("Diagram")
            pdf.diagram_text(sec["diagram"])
        if sec.get("body"):
            pdf.sub_title("Details")
            body = sec["body"].strip()
            if body.startswith("|") or body.startswith("[") or body.startswith("+"):
                pdf.diagram_text(body)
            else:
                for para in body.split("\n\n"):
                    para = para.strip()
                    if not para:
                        continue
                    if para.startswith("|") or para.startswith("["):
                        pdf.diagram_text(para)
                    else:
                        pdf.body_text(para)

    OUT_DIR.mkdir(parents=True, exist_ok=True)
    pdf.output(str(OUT_PDF))
    return OUT_PDF


def main() -> None:
    print("Generating LegalEase Blueprint documents...")
    docx_path = build_docx()
    print(f"  Word: {docx_path}")
    pdf_path = build_pdf()
    print(f"  PDF:  {pdf_path}")
    print("Done.")


if __name__ == "__main__":
    main()
