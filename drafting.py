"""
LegalEase.AI Automated Legal Drafting Module
=============================================
Generate professional legal documents using templates and LLM.
"""

import os
import json
from datetime import datetime
from pathlib import Path
from typing import Dict, Any, Optional

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Import LLM for AI-powered drafting
try:
    from llms import get_generator
    LLM_AVAILABLE = True
except ImportError:
    LLM_AVAILABLE = False


# Legal Document Templates 
TEMPLATES = {
    "LEGAL_NOTICE": """
LEGAL NOTICE
============

Date: {date}
Ref: {reference_number}

To,
{recipient_name}
{recipient_address}

Subject: Legal Notice for {subject}

Dear Sir/Madam,

Under instructions from and on behalf of my client, {client_name}, residing at {client_address}, I hereby serve upon you the following legal notice:

FACTS OF THE CASE:
{facts}

LEGAL GROUNDS:
{legal_grounds}

DEMAND:
In view of the above facts and circumstances, my client hereby demands that you:
{demands}

CONSEQUENCES OF NON-COMPLIANCE:
Please note that if you fail to comply with the above demands within {notice_period} days from the receipt of this notice, my client shall be constrained to initiate appropriate legal proceedings against you, civil and/or criminal, at your risk, cost and consequences.

You are hereby called upon to treat this notice as a statutory notice under the relevant provisions of law.

Yours faithfully,

{advocate_name}
Advocate
{advocate_enrollment}
{advocate_address}
{advocate_phone}

Copy to:
1. {client_name} (Client)
""",

    "AFFIDAVIT": """
AFFIDAVIT
=========

I, {deponent_name}, {deponent_relation} {deponent_father}, aged about {deponent_age} years, 
resident of {deponent_address}, do hereby solemnly affirm and state as under:

1. That I am the {deponent_capacity} in the above matter and am competent to swear this affidavit.

2. That I have personal knowledge of the facts stated herein.

{affidavit_content}

VERIFICATION:
I, {deponent_name}, the above-named deponent, do hereby verify that the contents of the above 
affidavit are true and correct to the best of my knowledge and belief, and nothing material 
has been concealed therefrom.

Verified at {verification_place} on this {date}.

DEPONENT

BEFORE ME:

Notary Public / Oath Commissioner
""",

    "CHARGESHEET": """
CHARGESHEET / FINAL REPORT
==========================

IN THE COURT OF {court_name}
{court_address}

Case No.: {case_number}
Under Sections: {sections}
Police Station: {police_station}
FIR No.: {fir_number}
Date of FIR: {fir_date}

STATE vs. {accused_name}

1. PARTICULARS OF ACCUSED:
   Name: {accused_name}
   Father's Name: {accused_father}
   Age: {accused_age}
   Address: {accused_address}
   Occupation: {accused_occupation}

2. BRIEF FACTS OF THE CASE:
   {case_facts}

3. INVESTIGATION SUMMARY:
   {investigation_summary}

4. EVIDENCE COLLECTED:
   Documentary Evidence:
   {documentary_evidence}
   
   Material Evidence:
   {material_evidence}

5. WITNESSES:
   {witnesses}

6. OPINION:
   Based on the investigation conducted, it is submitted that there is sufficient 
   evidence to prosecute the accused under the sections mentioned above.

7. PRAYER:
   It is therefore prayed that this Hon'ble Court may be pleased to take cognizance 
   of the offenses and issue process against the accused.

Date: {date}
Place: {place}

Investigating Officer
{io_name}
{io_rank}
{police_station}
""",

    "CONTRACT": """
CONTRACT AGREEMENT
==================

This Contract Agreement ("Agreement") is made and entered into on {date}

BETWEEN:

PARTY A (First Party):
Name: {party_a_name}
Address: {party_a_address}
PAN: {party_a_pan}
(Hereinafter referred to as "First Party")

AND

PARTY B (Second Party):
Name: {party_b_name}
Address: {party_b_address}
PAN: {party_b_pan}
(Hereinafter referred to as "Second Party")

RECITALS:
{recitals}

NOW THEREFORE, in consideration of the mutual covenants contained herein, the parties agree as follows:

1. DEFINITIONS
   {definitions}

2. SCOPE OF AGREEMENT
   {scope}

3. TERM AND DURATION
   This Agreement shall be effective from {start_date} and shall continue until {end_date}, 
   unless terminated earlier in accordance with the provisions hereof.

4. CONSIDERATION
   {consideration}

5. OBLIGATIONS OF FIRST PARTY
   {party_a_obligations}

6. OBLIGATIONS OF SECOND PARTY
   {party_b_obligations}

7. CONFIDENTIALITY
   Both parties agree to maintain confidentiality of all proprietary information 
   disclosed during the term of this Agreement.

8. TERMINATION
   {termination_clause}

9. DISPUTE RESOLUTION
   Any dispute arising out of this Agreement shall be resolved through arbitration 
   in accordance with the Arbitration and Conciliation Act, 1996. 
   The seat of arbitration shall be {arbitration_seat}.

10. GOVERNING LAW
    This Agreement shall be governed by the laws of India and the courts at 
    {jurisdiction} shall have exclusive jurisdiction.

11. GENERAL PROVISIONS
    {general_provisions}

IN WITNESS WHEREOF, the parties have executed this Agreement as of the date first written above.

FIRST PARTY:                          SECOND PARTY:
Signature: _______________            Signature: _______________
Name: {party_a_name}                  Name: {party_b_name}
Date: {date}                          Date: {date}

WITNESSES:
1. Name: _______________              2. Name: _______________
   Signature: _______________            Signature: _______________
   Address: _______________              Address: _______________
""",

    "BAIL_APPLICATION": """
BAIL APPLICATION
================

IN THE COURT OF {court_name}
{court_address}

Criminal Case No.: {case_number}
FIR No.: {fir_number}
Police Station: {police_station}
Under Sections: {sections}

IN THE MATTER OF:
{applicant_name} ... Applicant/Accused

VERSUS

State of {state} ... Respondent

APPLICATION FOR GRANT OF REGULAR/ANTICIPATORY BAIL

MOST RESPECTFULLY SHOWETH:

1. That the applicant is accused in the above-mentioned case registered at 
   {police_station} under sections {sections}.

2. BRIEF FACTS:
   {case_facts}

3. GROUNDS FOR BAIL:
   {bail_grounds}

4. That the applicant undertakes to:
   a) Appear before the Court as and when required
   b) Not tamper with evidence or influence witnesses
   c) Not leave the jurisdiction without permission
   d) Cooperate with the investigation

5. That the applicant is ready to furnish bail bond and surety as directed by this Hon'ble Court.

PRAYER:
In view of the above facts and circumstances, it is most respectfully prayed that this 
Hon'ble Court may be pleased to grant regular/anticipatory bail to the applicant in 
the interest of justice.

Date: {date}
Place: {place}

{applicant_name}
Through Counsel

{advocate_name}
Advocate
Enrollment No.: {advocate_enrollment}
"""
}


def generate_draft(template_name: str, context: Dict[str, Any], use_ai: bool = True) -> str:
    """
    Generate a legal document draft.
    
    Args:
        template_name: Name of the template (LEGAL_NOTICE, AFFIDAVIT, etc.)
        context: Dictionary with placeholder values
        use_ai: Whether to use AI to enhance the draft
    
    Returns:
        Generated document text
    """
    template_upper = template_name.upper().replace(" ", "_")
    
    if template_upper not in TEMPLATES:
        # Use AI to generate custom template
        if LLM_AVAILABLE and use_ai:
            generator = get_generator()
            prompt = f"""Generate a professional Indian legal document template for: {template_name}
            
Context provided:
{json.dumps(context, indent=2)}

Requirements:
1. Use proper legal formatting
2. Include all necessary sections
3. Use formal legal language
4. Follow Indian legal standards
5. Include relevant BNS/IPC sections where applicable

Generate the complete document:"""
            return generator.generate(prompt, temperature=0.3, max_tokens=2048)
        else:
            return f"Template '{template_name}' not found. Available templates: {', '.join(TEMPLATES.keys())}"
    
    # Get template
    template = TEMPLATES[template_upper]
    
    # Fill in default values for missing placeholders
    defaults = {
        "date": datetime.now().strftime("%d-%m-%Y"),
        "time": datetime.now().strftime("%H:%M"),
        "place": "[PLACE]",
        "state": "[STATE]",
        "district": "[DISTRICT]",
    }
    
    # Merge defaults with provided context
    full_context = {**defaults, **context}
    
    # Replace placeholders
    try:
        draft = template.format(**full_context)
    except KeyError as e:
        # If some placeholders are missing, replace them with [BLANK]
        import re
        draft = template
        for key in re.findall(r'\{(\w+)\}', template):
            if key in full_context:
                draft = draft.replace(f'{{{key}}}', str(full_context[key]))
            else:
                draft = draft.replace(f'{{{key}}}', f'[{key.upper()}]')
    
    # Optionally enhance with AI
    if LLM_AVAILABLE and use_ai and context.get("enhance_with_ai", False):
        generator = get_generator()
        enhance_prompt = f"""Review and enhance this legal document draft. 
Ensure it is:
1. Legally accurate
2. Properly formatted
3. Complete with all necessary sections
4. Using appropriate legal terminology

Original Draft:
{draft}

Enhanced Draft:"""
        draft = generator.generate(enhance_prompt, temperature=0.2, max_tokens=2048)
    
    return draft


def save_docx(text: str, output_path: Path) -> bool:
    """
    Save document text as a Word document.
    
    Args:
        text: Document text
        output_path: Path to save the document
    
    Returns:
        True if successful, False otherwise
    """
    if not DOCX_AVAILABLE:
        # Fallback to plain text
        output_path = Path(str(output_path).replace('.docx', '.txt'))
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w') as f:
            f.write(text)
        return True
    
    try:
        doc = Document()
        
        # Add title
        lines = text.strip().split('\n')
        if lines:
            title = doc.add_heading(lines[0].strip('=').strip(), 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content
        for line in lines[1:]:
            if line.strip('=').strip() == line.strip() and line.strip():
                # Check if it's a section header
                if line.strip().endswith(':') or line.strip().isupper():
                    doc.add_heading(line.strip(), level=1)
                else:
                    doc.add_paragraph(line)
            elif '=' in line:
                # Separator line, skip
                continue
        
        # Save document
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        return True
    
    except Exception as e:
        # Fallback to plain text
        output_path = Path(str(output_path).replace('.docx', '.txt'))
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, 'w') as f:
            f.write(text)
        return True


def get_available_templates() -> list:
    """Get list of available document templates."""
    return list(TEMPLATES.keys())


def get_template_fields(template_name: str) -> list:
    """Get required fields for a template."""
    import re
    template_upper = template_name.upper().replace(" ", "_")
    
    if template_upper not in TEMPLATES:
        return []
    
    template = TEMPLATES[template_upper]
    fields = re.findall(r'\{(\w+)\}', template)
    return list(set(fields))
