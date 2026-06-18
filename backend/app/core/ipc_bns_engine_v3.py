"""
IPC ↔ BNS Intelligence Engine V3 — deterministic database lookups only.
No LLM mapping. Unmapped sections return explicit verification message.
"""
from __future__ import annotations

import json
import re
import uuid
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from backend.app.core.database import connect_data_db
from backend.app.core.sql_compat import insert_or_replace

NOT_FOUND_MSG = "Official mapping not found. Manual legal verification required."

DATASET_VERSION = "2026.06.04-official"
DATASET_PATH = Path(__file__).resolve().parents[3] / "data" / "legal_conversion" / "ipc_bns_official.json"

_SECTION_PATTERNS = [
    re.compile(r"(?:IPC|I\.P\.C\.?)\s*(?:Section|Sec\.?|S\.?)?\s*(\d+[A-Z]?)", re.I),
    re.compile(r"(?:BNS|Bharatiya\s+Nyaya)\s*(?:Section|Sec\.?|S\.?)?\s*(\d+(?:\(\d+\))?)", re.I),
    re.compile(r"Section\s+(\d+[A-Z]?)\s+(?:of\s+)?(?:IPC|I\.P\.C\.?|Indian\s+Penal\s+Code)", re.I),
    re.compile(r"u/s\.?\s*(\d+[A-Z]?)\s*(?:IPC|I\.P\.C\.?|of\s+IPC)", re.I),
    re.compile(r"\b(?:IPC|I\.P\.C\.?)\s*(\d+[A-Z]?)\b", re.I),
    re.compile(r"\b(\d{1,4}[A-Z]?)\s*(?:IPC|I\.P\.C\.?)\b", re.I),
]

_BNS_PATTERN = re.compile(r"(?:BNS|Bharatiya\s+Nyaya)\s*(?:Section|Sec\.?|S\.?)?\s*(\d+(?:\(\d+\))?)", re.I)


def _utc() -> str:
    return datetime.now(timezone.utc).isoformat()


def normalize_ipc_key(raw: str) -> str:
    s = (raw or "").strip().upper()
    s = re.sub(r"^(?:IPC|SECTION|SEC\.?|S\.?)\s*", "", s, flags=re.I)
    s = s.strip()
    m = re.match(r"^(\d+[A-Z]?)", s)
    return m.group(1) if m else s


def normalize_bns_key(raw: str) -> str:
    s = (raw or "").strip()
    s = re.sub(r"^(?:BNS|SECTION|SEC\.?|S\.?)\s*", "", s, flags=re.I)
    s = s.replace(" ", "")
    m = re.match(r"^(\d+(?:\(\d+\))?)", s)
    return m.group(1) if m else s


def ensure_ipc_bns_schema() -> None:
    from backend.app.core.legal_conversion_engine import ensure_legal_conversion_schema

    ensure_legal_conversion_schema()


def _seed_dataset(conn) -> None:
    if not DATASET_PATH.is_file():
        try:
            from legal_tools import IPC_TO_BNS_MAP

            records = []
            for ipc_key, m in IPC_TO_BNS_MAP.items():
                records.append(
                    {
                        "ipc_key": ipc_key,
                        "bns_key": str(m["bns"]).replace(" ", ""),
                        "short_description": m["description"],
                        "offence_title": m["description"],
                    }
                )
            payload = {"records": records}
        except Exception:
            return
    else:
        payload = json.loads(DATASET_PATH.read_text(encoding="utf-8"))
    for r in payload.get("records") or []:
        params = (
            r.get("ipc_key"),
            r.get("bns_key", ""),
            r.get("ipc_section"),
            r.get("bns_section"),
            r.get("offence_title"),
            r.get("short_description"),
            r.get("punishment"),
            r.get("cognizable"),
            r.get("bailable"),
            r.get("compoundable"),
            r.get("court_jurisdiction"),
            r.get("mapping_status", "mapped"),
            r.get("change_notes"),
            r.get("legislative_notes"),
            r.get("source"),
            r.get("official_source"),
            r.get("gazette_reference"),
            r.get("dataset_version", DATASET_VERSION),
            r.get("last_updated"),
            r.get("keywords"),
            r.get("what_changed"),
            r.get("what_same"),
            r.get("punishment_changed"),
            r.get("procedure_changed"),
            r.get("new_definitions"),
            r.get("removed_definitions"),
            r.get("scope_expanded"),
            r.get("scope_reduced"),
        )
        insert_or_replace(
            conn,
            """
            INSERT OR REPLACE INTO ipc_bns_mappings (
                ipc_key, bns_key, ipc_section, bns_section, offence_title, short_description,
                punishment, cognizable, bailable, compoundable, court_jurisdiction, mapping_status,
                change_notes, legislative_notes, source, official_source, gazette_reference,
                dataset_version, last_updated, keywords, what_changed, what_same,
                punishment_changed, procedure_changed, new_definitions, removed_definitions,
                scope_expanded, scope_reduced
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            """,
            """
            INSERT INTO ipc_bns_mappings (
                ipc_key, bns_key, ipc_section, bns_section, offence_title, short_description,
                punishment, cognizable, bailable, compoundable, court_jurisdiction, mapping_status,
                change_notes, legislative_notes, source, official_source, gazette_reference,
                dataset_version, last_updated, keywords, what_changed, what_same,
                punishment_changed, procedure_changed, new_definitions, removed_definitions,
                scope_expanded, scope_reduced
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)
            ON CONFLICT (ipc_key) DO UPDATE SET
                bns_key = EXCLUDED.bns_key,
                ipc_section = EXCLUDED.ipc_section,
                bns_section = EXCLUDED.bns_section,
                offence_title = EXCLUDED.offence_title,
                short_description = EXCLUDED.short_description,
                punishment = EXCLUDED.punishment,
                cognizable = EXCLUDED.cognizable,
                bailable = EXCLUDED.bailable,
                compoundable = EXCLUDED.compoundable,
                court_jurisdiction = EXCLUDED.court_jurisdiction,
                mapping_status = EXCLUDED.mapping_status,
                change_notes = EXCLUDED.change_notes,
                legislative_notes = EXCLUDED.legislative_notes,
                source = EXCLUDED.source,
                official_source = EXCLUDED.official_source,
                gazette_reference = EXCLUDED.gazette_reference,
                dataset_version = EXCLUDED.dataset_version,
                last_updated = EXCLUDED.last_updated,
                keywords = EXCLUDED.keywords,
                what_changed = EXCLUDED.what_changed,
                what_same = EXCLUDED.what_same,
                punishment_changed = EXCLUDED.punishment_changed,
                procedure_changed = EXCLUDED.procedure_changed,
                new_definitions = EXCLUDED.new_definitions,
                removed_definitions = EXCLUDED.removed_definitions,
                scope_expanded = EXCLUDED.scope_expanded,
                scope_reduced = EXCLUDED.scope_reduced
            """,
            params,
        )
    conn.commit()


def _row_to_dict(row: Tuple) -> Dict[str, Any]:
    cols = [
        "ipc_key",
        "bns_key",
        "ipc_section",
        "bns_section",
        "offence_title",
        "short_description",
        "punishment",
        "cognizable",
        "bailable",
        "compoundable",
        "court_jurisdiction",
        "mapping_status",
        "change_notes",
        "legislative_notes",
        "source",
        "official_source",
        "gazette_reference",
        "dataset_version",
        "last_updated",
        "keywords",
        "what_changed",
        "what_same",
        "punishment_changed",
        "procedure_changed",
        "new_definitions",
        "removed_definitions",
        "scope_expanded",
        "scope_reduced",
    ]
    return dict(zip(cols, row))


def _fetch_row_ipc(conn, ipc_key: str) -> Optional[Dict[str, Any]]:
    row = conn.execute(
        "SELECT * FROM ipc_bns_mappings WHERE ipc_key = ?",
        (ipc_key,),
    ).fetchone()
    return _row_to_dict(row) if row else None


def _fetch_row_bns(conn, bns_key: str) -> Optional[Dict[str, Any]]:
    row = conn.execute(
        "SELECT * FROM ipc_bns_mappings WHERE bns_key = ? OR bns_key LIKE ? LIMIT 1",
        (bns_key, f"{bns_key}%"),
    ).fetchone()
    return _row_to_dict(row) if row else None


def _not_found(ipc_key: str = "", bns_key: str = "") -> Dict[str, Any]:
    return {
        "status": "not_found",
        "found": False,
        "ipc_key": ipc_key or None,
        "bns_key": bns_key or None,
        "ipc_section": f"IPC Section {ipc_key}" if ipc_key else None,
        "bns_section": f"BNS Section {bns_key}" if bns_key else None,
        "message": NOT_FOUND_MSG,
        "source": "LegalEase IPC-BNS dataset",
        "dataset_version": DATASET_VERSION,
        "confidence": 0,
    }


def _found_record(rec: Dict[str, Any]) -> Dict[str, Any]:
    return {
        **rec,
        "status": "mapped",
        "found": True,
        "confidence": 100,
        "message": None,
    }


def audit_log(
    user_id: str,
    action: str,
    query: str = "",
    result_summary: str = "",
    matter_id: str = "",
) -> str:
    ensure_ipc_bns_schema()
    aid = str(uuid.uuid4())
    conn = connect_data_db()
    conn.execute(
        """
        INSERT INTO ipc_bns_audit (audit_id, user_id, action, query, result_summary, dataset_version, matter_id, created_at)
        VALUES (?,?,?,?,?,?,?,?)
        """,
        (aid, str(user_id), action, query[:2000], result_summary[:4000], DATASET_VERSION, matter_id or "", _utc()),
    )
    conn.commit()
    conn.close()
    return aid


def lookup_ipc(ipc_section: str, *, user_id: str = "", matter_id: str = "") -> Dict[str, Any]:
    from backend.app.core.legal_conversion_engine import convert_section, to_ipc_bns_v3_shape

    ensure_ipc_bns_schema()
    raw = convert_section("ipc_bns", ipc_section, direction="forward", user_id=user_id, matter_id=matter_id)
    return to_ipc_bns_v3_shape(raw)


def lookup_bns(bns_section: str, *, user_id: str = "", matter_id: str = "") -> Dict[str, Any]:
    from backend.app.core.legal_conversion_engine import convert_section, to_ipc_bns_v3_shape

    ensure_ipc_bns_schema()
    raw = convert_section("ipc_bns", bns_section, direction="reverse", user_id=user_id, matter_id=matter_id)
    return to_ipc_bns_v3_shape(raw)


def search_mappings(q: str, *, limit: int = 25) -> Dict[str, Any]:
    from backend.app.core.legal_conversion_engine import search_mappings as lc_search
    from backend.app.core.legal_conversion_engine import to_ipc_bns_v3_shape

    ensure_ipc_bns_schema()
    raw = lc_search("ipc_bns", q, limit=limit)
    return {
        "results": [to_ipc_bns_v3_shape(r) for r in raw.get("results") or []],
        "count": raw.get("count", 0),
        "dataset_version": DATASET_VERSION,
    }


def compare_sections(ipc_section: str) -> Dict[str, Any]:
    rec = lookup_ipc(ipc_section)
    if not rec.get("found"):
        return rec
    return {
        "ipc": rec,
        "bns": rec,
        "comparison": {
            "what_changed": rec.get("what_changed"),
            "what_same": rec.get("what_same"),
            "punishment_changed": rec.get("punishment_changed"),
            "procedure_changed": rec.get("procedure_changed"),
            "new_definitions": rec.get("new_definitions"),
            "removed_definitions": rec.get("removed_definitions"),
            "scope_expanded": rec.get("scope_expanded"),
            "scope_reduced": rec.get("scope_reduced"),
            "change_notes": rec.get("change_notes"),
        },
        "dataset_version": DATASET_VERSION,
        "source": rec.get("source"),
        "official_source": rec.get("official_source"),
        "gazette_reference": rec.get("gazette_reference"),
        "last_updated": rec.get("last_updated"),
    }


def extract_ipc_sections(text: str) -> List[str]:
    found: List[str] = []
    seen = set()
    for pat in _SECTION_PATTERNS:
        for m in pat.finditer(text or ""):
            key = normalize_ipc_key(m.group(1))
            if key and key not in seen and re.match(r"^\d+[A-Z]?$", key):
                seen.add(key)
                found.append(key)
    return found


def bulk_convert_ipc(sections: List[str], *, user_id: str = "", matter_id: str = "") -> Dict[str, Any]:
    results = []
    mapped = 0
    for s in sections:
        r = lookup_ipc(s, user_id=user_id, matter_id=matter_id)
        results.append(r)
        if r.get("found"):
            mapped += 1
    if user_id:
        audit_log(user_id, "bulk_convert", ",".join(sections[:50]), f"{mapped}/{len(sections)} mapped", matter_id)
    return {
        "results": results,
        "total": len(results),
        "mapped_count": mapped,
        "unmapped_count": len(results) - mapped,
        "dataset_version": DATASET_VERSION,
    }


def extract_and_convert_document(text: str, *, user_id: str = "", matter_id: str = "") -> Dict[str, Any]:
    keys = extract_ipc_sections(text)
    bulk = bulk_convert_ipc(keys, user_id=user_id, matter_id=matter_id)
    return {"sections_detected": keys, **bulk}


def dataset_meta() -> Dict[str, Any]:
    from backend.app.core.legal_conversion_engine import dataset_meta as lc_meta

    ensure_ipc_bns_schema()
    meta = lc_meta()
    ipc_pair = next((p for p in meta.get("pairs") or [] if p.get("pair_type") == "ipc_bns"), {})
    return {
        **meta,
        "record_count": ipc_pair.get("record_count", 0),
        "source": "Official IPC↔BNS mapping dataset (ipc_bns_official.json)",
        "official_source": ipc_pair.get("source_file"),
        "gazette_reference": "Bharatiya Nyaya Sanhita, 2023 (Act 45 of 2023)",
        "last_updated": meta.get("last_updated"),
        "deterministic": True,
        "ai_mapping": False,
    }


def list_categories() -> List[str]:
    return [
        "murder",
        "theft",
        "robbery",
        "assault",
        "sexual_offenses",
        "cheating",
        "forgery",
        "defamation",
        "kidnapping",
        "dowry",
        "trespass",
        "criminal_intimidation",
    ]


def sections_by_category(category: str) -> List[Dict[str, Any]]:
    from legal_tools import get_bns_by_category

    out = []
    for item in get_bns_by_category(category):
        ipc_num = re.sub(r"\D", "", item.get("ipc_section", ""))
        if ipc_num:
            r = lookup_ipc(ipc_num)
            if r.get("found"):
                out.append(r)
            else:
                out.append({**item, "status": "not_found", "message": NOT_FOUND_MSG})
    return out


def extract_text_from_upload(filename: str, data: bytes) -> str:
    name = (filename or "").lower()
    if name.endswith(".txt"):
        return data.decode("utf-8", errors="replace")
    if name.endswith(".pdf"):
        from PyPDF2 import PdfReader

        reader = PdfReader(BytesIO(data))
        return "\n".join((p.extract_text() or "") for p in reader.pages)
    if name.endswith(".docx"):
        from docx import Document

        doc = Document(BytesIO(data))
        return "\n".join(p.text for p in doc.paragraphs)
    return data.decode("utf-8", errors="replace")


def build_conversion_report(
    *,
    case_name: str,
    conversions: List[Dict[str, Any]],
    generated_by: str = "",
) -> Dict[str, Any]:
    return {
        "case_name": case_name or "IPC-BNS Migration Report",
        "date": datetime.now(timezone.utc).strftime("%d %B %Y"),
        "sections_detected": len(conversions),
        "conversions": conversions,
        "dataset_version": DATASET_VERSION,
        "generated_by": generated_by,
        "disclaimer": "Mappings from verified dataset only. Manual verification required for unmapped sections.",
    }


def report_to_pdf_bytes(report: Dict[str, Any]) -> bytes:
    from fpdf import FPDF

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "IPC - BNS Migration Report", ln=True)
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 8, f"Case: {report.get('case_name', '')}", ln=True)
    pdf.cell(0, 8, f"Date: {report.get('date', '')}", ln=True)
    pdf.cell(0, 8, f"Dataset: v{report.get('dataset_version', '')}", ln=True)
    pdf.ln(4)
    pdf.set_font("Helvetica", "B", 11)
    pdf.cell(40, 8, "IPC", border=1)
    pdf.cell(40, 8, "BNS", border=1)
    pdf.cell(110, 8, "Notes", border=1, ln=True)
    pdf.set_font("Helvetica", "", 9)
    for row in report.get("conversions") or []:
        ipc = str(row.get("ipc_section") or row.get("ipc_key") or "-")[:18]
        bns = str(row.get("bns_section") or row.get("bns_key") or row.get("message") or "-")[:18]
        note = str(row.get("change_notes") or row.get("short_description") or row.get("message") or "")[:60]
        pdf.cell(40, 7, ipc, border=1)
        pdf.cell(40, 7, bns, border=1)
        pdf.cell(110, 7, note, border=1, ln=True)
    pdf.ln(6)
    pdf.set_font("Helvetica", "I", 8)
    pdf.multi_cell(0, 5, str(report.get("disclaimer", "")))
    return bytes(pdf.output(dest="S"))


def report_to_docx_bytes(report: Dict[str, Any]) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading("IPC - BNS Migration Report", 0)
    doc.add_paragraph(f"Case: {report.get('case_name', '')}")
    doc.add_paragraph(f"Date: {report.get('date', '')}")
    doc.add_paragraph(f"Dataset version: {report.get('dataset_version', '')}")
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "IPC"
    hdr[1].text = "BNS"
    hdr[2].text = "Notes"
    for row in report.get("conversions") or []:
        cells = table.add_row().cells
        cells[0].text = str(row.get("ipc_section") or row.get("ipc_key") or "")
        cells[1].text = str(row.get("bns_section") or row.get("bns_key") or row.get("message") or "")
        cells[2].text = str(row.get("change_notes") or row.get("short_description") or row.get("message") or "")
    doc.add_paragraph(str(report.get("disclaimer", "")))
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def matter_migration_impact(user_id: str, matter_id: str) -> Dict[str, Any]:
    from backend.app.core.matter_repo import get_matter
    from backend.app.core.drafting_workspace import list_documents

    matter = get_matter(user_id, matter_id)
    if not matter:
        return {"error": "Matter not found"}
    corpus_parts = [
        matter.get("matter_name") or "",
        matter.get("case_summary") or "",
        matter.get("notes") or "",
    ]
    docs = [d for d in list_documents(user_id, limit=100) if d.get("matter_id") == matter_id]
    for d in docs[:20]:
        corpus_parts.append(d.get("content") or "")
        corpus_parts.append(d.get("title") or "")
    text = "\n".join(corpus_parts)
    conv = extract_and_convert_document(text, user_id=user_id, matter_id=matter_id)
    affected = [r for r in conv.get("results") or [] if r.get("found")]
    unmapped = [r for r in conv.get("results") or [] if not r.get("found")]
    memo_lines = [
        f"Matter: {matter.get('matter_name', matter_id)}",
        f"IPC sections detected: {len(conv.get('sections_detected') or [])}",
        f"Official mappings: {len(affected)}",
        f"Requires manual verification: {len(unmapped)}",
    ]
    for r in affected[:15]:
        memo_lines.append(f"  - {r.get('ipc_section')} -> {r.get('bns_section')}: {r.get('offence_title')}")
    for r in unmapped[:10]:
        memo_lines.append(f"  - {r.get('ipc_section')}: {NOT_FOUND_MSG}")
    return {
        "matter_id": matter_id,
        "migration_impact": {
            "sections_detected": conv.get("sections_detected"),
            "mapped": affected,
            "unmapped": unmapped,
            "mapped_count": len(affected),
            "unmapped_count": len(unmapped),
        },
        "migration_memo": "\n".join(memo_lines),
        "dataset_version": DATASET_VERSION,
    }
