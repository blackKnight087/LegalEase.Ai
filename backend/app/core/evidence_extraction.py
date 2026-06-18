"""Multi-format evidence file text extraction for the Evidence Intelligence Center."""
from __future__ import annotations

import email
import hashlib
import io
import logging
import zipfile
from datetime import datetime, timezone
from email import policy
from pathlib import Path
from typing import Any, Dict, List, Tuple

logger = logging.getLogger("legalease.evidence_extract")

_SUPPORTED = {
    ".pdf", ".docx", ".doc", ".png", ".jpg", ".jpeg", ".webp", ".gif",
    ".tif", ".tiff", ".xlsx", ".xls", ".csv", ".txt", ".eml", ".msg", ".zip",
}
_IMAGE_EXT = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tif", ".tiff"}


def file_hash(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


def _utc_iso(ts: float | None = None) -> str:
    if ts is None:
        return datetime.now(timezone.utc).isoformat()
    return datetime.fromtimestamp(ts, tz=timezone.utc).isoformat()


def extract_metadata(filename: str, data: bytes) -> Dict[str, Any]:
    """Best-effort file metadata (author, dates, type, hash)."""
    name = filename or "evidence"
    ext = Path(name).suffix.lower()
    meta: Dict[str, Any] = {
        "filename": name,
        "file_type": ext.lstrip(".") or "unknown",
        "size_bytes": len(data),
        "sha256": file_hash(data),
        "author": "",
        "created_at": "",
        "modified_at": "",
    }
    if ext == ".pdf" and data[:4] == b"%PDF":
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(data))
            info = reader.metadata or {}
            meta["author"] = str(getattr(info, "author", "") or info.get("/Author", "") or "")
            meta["page_count"] = len(reader.pages)
        except Exception as exc:
            logger.debug("PDF metadata: %s", exc)
    elif ext == ".docx":
        try:
            from docx import Document

            doc = Document(io.BytesIO(data))
            cp = doc.core_properties
            meta["author"] = (cp.author or cp.last_modified_by or "")[:120]
            if cp.created:
                meta["created_at"] = cp.created.isoformat()
            if cp.modified:
                meta["modified_at"] = cp.modified.isoformat()
        except Exception as exc:
            logger.debug("DOCX metadata: %s", exc)
    elif ext == ".eml":
        try:
            msg = email.message_from_bytes(data, policy=policy.default)
            meta["author"] = (msg.get("From") or "")[:200]
            meta["created_at"] = (msg.get("Date") or "")[:80]
            meta["subject"] = (msg.get("Subject") or "")[:200]
        except Exception as exc:
            logger.debug("EML metadata: %s", exc)
    return meta


def _extract_xlsx(data: bytes) -> str:
    try:
        import openpyxl

        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: List[str] = []
        for sheet in wb.worksheets:
            for row in sheet.iter_rows(max_row=500, values_only=True):
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append("\t".join(cells))
        return "\n".join(parts)[:50000]
    except ImportError:
        pass
    except Exception as exc:
        logger.debug("openpyxl failed: %s", exc)
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            shared: List[str] = []
            if "xl/sharedStrings.xml" in zf.namelist():
                import xml.etree.ElementTree as ET

                root = ET.fromstring(zf.read("xl/sharedStrings.xml"))
                ns = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
                for si in root.findall(".//m:si", ns):
                    t = "".join(n.text or "" for n in si.findall(".//m:t", ns))
                    shared.append(t)
            sheet_name = next((n for n in zf.namelist() if n.startswith("xl/worksheets/sheet")), "")
            if sheet_name:
                import xml.etree.ElementTree as ET

                root = ET.fromstring(zf.read(sheet_name))
                ns = {"m": "http://schemas.openxmlformats.org/spreadsheetml/2006/main"}
                rows: List[str] = []
                for row in root.findall(".//m:row", ns):
                    cells = []
                    for c in row.findall("m:c", ns):
                        ref = c.get("t")
                        v = c.find("m:v", ns)
                        if v is not None and v.text:
                            if ref == "s":
                                idx = int(v.text)
                                cells.append(shared[idx] if idx < len(shared) else v.text)
                            else:
                                cells.append(v.text)
                    if cells:
                        rows.append("\t".join(cells))
                return "\n".join(rows)[:50000]
    except Exception as exc:
        logger.debug("xlsx xml fallback: %s", exc)
    return ""


def _extract_eml(data: bytes) -> str:
    msg = email.message_from_bytes(data, policy=policy.default)
    parts: List[str] = []
    hdr_from = msg.get("From") or ""
    hdr_to = msg.get("To") or ""
    hdr_sub = msg.get("Subject") or ""
    hdr_date = msg.get("Date") or ""
    if hdr_from:
        parts.append(f"From: {hdr_from}")
    if hdr_to:
        parts.append(f"To: {hdr_to}")
    if hdr_sub:
        parts.append(f"Subject: {hdr_sub}")
    if hdr_date:
        parts.append(f"Date: {hdr_date}")
    body = ""
    if msg.is_multipart():
        for part in msg.walk():
            if part.get_content_type() == "text/plain":
                body = part.get_content() or ""
                break
    else:
        body = msg.get_content() or ""
    if body:
        parts.append(str(body))
    return "\n".join(parts).strip()[:50000]


def _extract_single(filename: str, data: bytes) -> Tuple[str, str]:
    """Returns (text, extraction_method)."""
    name = (filename or "upload").lower()
    ext = Path(name).suffix.lower()

    if ext in (".txt", ".md", ".csv") or (not ext and data[:20].isascii()):
        return data.decode("utf-8", errors="replace").strip()[:50000], "plain_text"

    if ext == ".eml" or b"From:" in data[:200]:
        return _extract_eml(data), "email_eml"

    if ext == ".msg":
        try:
            import extract_msg

            with io.BytesIO(data) as buf:
                msg = extract_msg.Message(buf)
                parts = [f"From: {msg.sender or ''}", f"Subject: {msg.subject or ''}", msg.body or ""]
                return "\n".join(p for p in parts if p).strip()[:50000], "email_msg"
        except Exception:
            return data.decode("utf-8", errors="replace").strip()[:50000], "msg_fallback"

    if ext in _IMAGE_EXT or data[:8] == b"\x89PNG\r\n\x1a\n" or data[:3] == b"\xff\xd8\xff":
        try:
            from ocr_engine import extract_text_from_image_bytes

            text, method = extract_text_from_image_bytes(data, filename or "evidence.png")
            return (text or "").strip()[:50000], method or "ocr_image"
        except Exception as exc:
            logger.warning("Image OCR failed: %s", exc)
            return "", "ocr_failed"

    if ext == ".pdf" or data[:4] == b"%PDF":
        import tempfile

        with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
            tmp.write(data)
            tmp_path = Path(tmp.name)
        try:
            from backend.app.core.pdf_extraction import extract_pdf_production

            text, method = extract_pdf_production(tmp_path, allow_ocr=True)
            return (text or "").strip()[:50000], method or "pdf"
        except Exception:
            from PyPDF2 import PdfReader

            reader = PdfReader(io.BytesIO(data))
            text = "\n".join((p.extract_text() or "").strip() for p in reader.pages)
            return text.strip()[:50000], "pdf_native"
        finally:
            try:
                tmp_path.unlink(missing_ok=True)
            except OSError:
                pass

    if ext in (".docx", ".doc"):
        try:
            from docx import Document

            doc = Document(io.BytesIO(data))
            text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
            return text.strip()[:50000], "docx"
        except Exception:
            from backend.app.core.crm_document_extract import extract_crm_upload_text

            return extract_crm_upload_text(data, filename), "docx_fallback"

    if ext in (".xlsx", ".xls"):
        text = _extract_xlsx(data)
        return text[:50000], "xlsx" if text else "xlsx_empty"

    if ext == ".zip":
        return _extract_zip(data), "zip_archive"

    return data.decode("utf-8", errors="replace").strip()[:50000], "binary_fallback"


def _extract_zip(data: bytes, *, depth: int = 0) -> str:
    if depth > 2:
        return ""
    parts: List[str] = []
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as zf:
            for info in zf.infolist():
                if info.is_dir():
                    continue
                inner_name = info.filename
                inner_ext = Path(inner_name).suffix.lower()
                if inner_ext not in _SUPPORTED and inner_ext not in _IMAGE_EXT:
                    continue
                try:
                    inner_data = zf.read(info)
                    text, _ = _extract_single(inner_name, inner_data)
                    if text.strip():
                        parts.append(f"--- {inner_name} ---\n{text.strip()}")
                except Exception as exc:
                    logger.debug("ZIP inner %s: %s", inner_name, exc)
    except Exception as exc:
        logger.warning("ZIP extract failed: %s", exc)
    return "\n\n".join(parts)[:50000]


def extract_evidence_file(filename: str, data: bytes) -> Dict[str, Any]:
    """Full extraction result with text, metadata, and method."""
    name = filename or "evidence"
    ext = Path(name).suffix.lower()
    if ext == ".zip":
        text, method = _extract_zip(data), "zip_archive"
    else:
        text, method = _extract_single(name, data)
    metadata = extract_metadata(name, data)
    metadata["extraction_method"] = method
    return {
        "text": text,
        "extraction_method": method,
        "metadata": metadata,
        "supported": ext in _SUPPORTED or ext in _IMAGE_EXT or ext == "",
    }


def supported_extensions() -> List[str]:
    return sorted(_SUPPORTED | _IMAGE_EXT)
