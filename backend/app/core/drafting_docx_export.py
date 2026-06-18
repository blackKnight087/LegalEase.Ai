"""Export legal HTML to formatted DOCX (Word round-trip quality)."""
from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from html.parser import HTMLParser
from typing import Any, Dict, List, Tuple

from backend.app.core.drafting_document_format import normalize_content_for_storage, sanitize_html


class _DocxBlockParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self.blocks: List[Dict[str, Any]] = []
        self._table: List[List[str]] = []
        self._row: List[str] = []
        self._cell = ""
        self._in_table = False
        self._in_cell = False
        self._tag_stack: List[str] = []
        self._list_items: List[str] = []
        self._in_li = False
        self._li_buf = ""

    def handle_starttag(self, tag: str, attrs: List) -> None:
        t = tag.lower()
        self._tag_stack.append(t)
        if t == "table":
            self._in_table = True
            self._table = []
        elif t == "tr" and self._in_table:
            self._row = []
        elif t in ("td", "th") and self._in_table:
            self._in_cell = True
            self._cell = ""
        elif t in ("h1", "h2", "h3"):
            self.blocks.append({"type": "heading", "level": int(t[1]), "text": ""})
        elif t == "p" and not self._in_table:
            self.blocks.append({"type": "paragraph", "text": "", "align": ""})
        elif t == "li":
            self._in_li = True
            self._li_buf = ""
        elif t == "hr":
            self.blocks.append({"type": "page_break"})
        elif t == "strong" or t == "b":
            pass

    def handle_endtag(self, tag: str) -> None:
        t = tag.lower()
        if t in ("td", "th") and self._in_cell:
            self._row.append(self._cell.strip())
            self._in_cell = False
        elif t == "tr" and self._in_table and self._row:
            self._table.append(self._row)
            self._row = []
        elif t == "table" and self._in_table:
            self.blocks.append({"type": "table", "rows": self._table})
            self._in_table = False
        elif t in ("h1", "h2", "h3", "p") and self.blocks and "text" in self.blocks[-1]:
            self.blocks[-1]["text"] = self.blocks[-1].get("text", "").strip()
        elif t == "li" and self._in_li:
            self._list_items.append(self._li_buf.strip())
            self._in_li = False
            if self._tag_stack.count("ul") or self._tag_stack.count("ol"):
                pass
        elif t in ("ul", "ol") and self._list_items:
            self.blocks.append({"type": "list", "items": list(self._list_items)})
            self._list_items = []

    def handle_data(self, data: str) -> None:
        if self._in_cell:
            self._cell += data
        elif self._in_li:
            self._li_buf += data
        elif self.blocks and not self._in_table:
            last = self.blocks[-1]
            if "text" in last:
                last["text"] = (last.get("text") or "") + data


def html_to_docx_bytes(
    html: str,
    *,
    title: str = "Legal Document",
    firm_name: str = "LegalEase",
) -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt

    html = sanitize_html(html)
    parser = _DocxBlockParser()
    try:
        parser.feed(html)
        parser.close()
    except Exception:
        parser.blocks = [{"type": "paragraph", "text": re.sub(r"<[^>]+>", "", html)}]

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover.add_run(firm_name)
    run.bold = True
    run.font.size = Pt(14)

    tpara = doc.add_paragraph()
    tpara.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tpara.add_run(title)
    tr.bold = True
    tr.font.size = Pt(18)

    doc.add_paragraph()

    for block in parser.blocks:
        btype = block.get("type")
        if btype == "page_break":
            doc.add_page_break()
        elif btype == "heading":
            lvl = int(block.get("level") or 2)
            text = (block.get("text") or "").strip()
            if text:
                doc.add_heading(text, level=min(lvl, 3))
        elif btype == "paragraph":
            text = (block.get("text") or "").strip()
            if text:
                p = doc.add_paragraph(text)
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(6)
        elif btype == "list":
            for item in block.get("items") or []:
                if item:
                    doc.add_paragraph(item, style="List Bullet")
        elif btype == "table":
            rows = block.get("rows") or []
            if not rows:
                continue
            cols = max(len(r) for r in rows)
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for ri, row in enumerate(rows):
                for ci in range(cols):
                    cell_text = row[ci] if ci < len(row) else ""
                    table.rows[ri].cells[ci].text = str(cell_text)[:200]

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def export_docx_document(
    content: str,
    *,
    title: str = "Legal Document",
    content_format: str = "html",
    firm_name: str = "LegalEase",
) -> Tuple[bytes, str, str]:
    html, _ = normalize_content_for_storage(content, content_format)
    safe = re.sub(r"[^\w\s-]", "", title or "document")[:60].strip() or "document"
    ts = datetime.now(timezone.utc).strftime("%Y%m%d")
    try:
        data = html_to_docx_bytes(html, title=title, firm_name=firm_name)
        return data, f"{safe}_{ts}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except ImportError:
        from backend.app.core.report_export import export_report_bytes
        from backend.app.core.drafting_document_format import html_to_plain_text

        return export_report_bytes(html_to_plain_text(html), title=title, fmt="docx")
