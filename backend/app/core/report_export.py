"""
Export Jurisprudence / research reports as DOCX or PDF.
"""
from __future__ import annotations

import io
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Tuple


def redact_for_client(content: str) -> str:
    """Remove internal KB markers, verification blocks, and obvious PII for client memos."""
    text = (content or "").strip()
    if not text:
        return text
    text = re.sub(r"\[KB-\d+\]", "", text, flags=re.I)
    text = re.sub(r"\[WEB-\d+\]", "", text, flags=re.I)
    text = re.sub(r"## Citation Verification[\s\S]*?(?=\n## |\Z)", "", text, flags=re.I)
    text = re.sub(
        r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b",
        "[email redacted]",
        text,
    )
    text = re.sub(r"\b(?:\+91[\s-]?)?[6-9]\d{9}\b", "[phone redacted]", text)
    text = re.sub(r"\b\d{4}[\s-]?\d{4}[\s-]?\d{4}\b", "[id redacted]", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _strip_md(text: str) -> str:
    t = text or ""
    t = re.sub(r"^#+\s*", "", t, flags=re.M)
    t = re.sub(r"\*\*([^*]+)\*\*", r"\1", t)
    t = re.sub(r"\*([^*]+)\*", r"\1", t)
    t = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", t)
    return t.strip()


def export_report_bytes(
    content: str,
    title: str = "LegalEase Research Report",
    fmt: str = "docx",
    *,
    client_safe: bool = False,
) -> Tuple[bytes, str, str]:
    """
    Returns (file_bytes, filename, media_type).
    fmt: docx | pdf | md
    """
    safe_title = re.sub(r"[^\w\s-]", "", title or "report")[:60].strip() or "report"
    if client_safe:
        safe_title = f"Client Memo - {safe_title}"[:72]
    ts = datetime.now(timezone.utc).strftime("%Y%m%d")
    body = redact_for_client(content) if client_safe else (content or "").strip()

    if fmt == "md":
        data = body.encode("utf-8")
        return data, f"{safe_title}_{ts}.md", "text/markdown"

    if fmt == "docx":
        return _export_docx(body, safe_title, ts)

    if fmt == "pdf":
        return _export_pdf(body, safe_title, ts)

    raise ValueError(f"Unsupported format: {fmt}")


def _export_docx(body: str, safe_title: str, ts: str) -> Tuple[bytes, str, str]:
    buf = io.BytesIO()
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        heading = doc.add_heading(safe_title, 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}")

        for line in body.splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("## "):
                doc.add_heading(_strip_md(stripped), level=2)
            elif stripped.startswith("### "):
                doc.add_heading(_strip_md(stripped), level=3)
            elif stripped.startswith("- "):
                doc.add_paragraph(_strip_md(stripped[2:]), style="List Bullet")
            else:
                doc.add_paragraph(_strip_md(stripped))

        doc.save(buf)
        buf.seek(0)
        return (
            buf.read(),
            f"{safe_title}_{ts}.docx",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except ImportError:
        data = body.encode("utf-8")
        return data, f"{safe_title}_{ts}.txt", "text/plain"


def _export_pdf(body: str, safe_title: str, ts: str) -> Tuple[bytes, str, str]:
    buf = io.BytesIO()
    try:
        import fitz  # PyMuPDF

        doc = fitz.open()
        page = doc.new_page(width=595, height=842)  # A4
        y = 50
        margin = 50
        width = 595 - 2 * margin

        def write_line(text: str, size: int = 11, bold: bool = False):
            nonlocal y, page
            if y > 780:
                page = doc.new_page(width=595, height=842)
                y = 50
            page.insert_text(
                (margin, y),
                text[:500],
                fontsize=size,
                fontname="helv",
            )
            y += size + 6

        write_line(safe_title, size=16)
        y += 4
        write_line(f"Generated {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M UTC')}", size=9)
        y += 8

        for line in body.splitlines():
            stripped = line.strip()
            if not stripped:
                y += 4
                continue
            if stripped.startswith("## "):
                y += 6
                write_line(_strip_md(stripped), size=13)
            elif stripped.startswith("### "):
                write_line(_strip_md(stripped), size=12)
            else:
                # Wrap long paragraphs simply
                plain = _strip_md(stripped)
                while plain:
                    chunk = plain[:90]
                    last_space = chunk.rfind(" ")
                    if len(plain) > 90 and last_space > 40:
                        chunk = plain[:last_space]
                        plain = plain[last_space + 1 :]
                    else:
                        plain = plain[90:] if len(plain) > 90 else ""
                    write_line(chunk if chunk else stripped[:90])

        doc.save(buf)
        doc.close()
        buf.seek(0)
        return buf.read(), f"{safe_title}_{ts}.pdf", "application/pdf"
    except Exception:
        # Fallback to plain text bytes
        data = body.encode("utf-8")
        return data, f"{safe_title}_{ts}.txt", "text/plain"
